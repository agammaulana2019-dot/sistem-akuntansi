import streamlit as st
import pandas as pd
from datetime import datetime
import io

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Konfigurasi halaman
st.set_page_config(
    page_title="Sistem Akuntansi Pro",
    page_icon="💰",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Custom untuk styling
st.markdown("""
    <style>
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    .stApp {
        background: linear-gradient(to bottom right, #f8f9fa, #e9ecef);
    }
    div[data-testid="stMetricValue"] {
        font-size: 28px;
        font-weight: bold;
        color: #2c3e50;
    }
    .css-1d391kg, .css-18e3th9 {
        padding-top: 2rem;
    }
    h1 {
        color: #2c3e50;
        font-weight: 800;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    h2 {
        color: #34495e;
        font-weight: 700;
        border-bottom: 3px solid #667eea;
        padding-bottom: 0.5rem;
    }
    h3 {
        color: #5a67d8;
    }
    .stButton>button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        border: none;
        padding: 0.5rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
    }
    div[data-testid="stDataFrame"] {
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .stSelectbox, .stDateInput, .stNumberInput, .stTextInput {
        border-radius: 8px;
    }
    div.stAlert {
        border-radius: 10px;
        border-left: 5px solid #667eea;
    }
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(180deg, #667eea 0%, #764ba2 100%);
    }
    </style>
""", unsafe_allow_html=True)

# Inisialisasi session state
if 'transaksi' not in st.session_state:
    st.session_state.transaksi = []

# Daftar akun standar dengan emoji
AKUN_HARTA = ['💵 Kas', '🏦 Bank', '📝 Piutang Usaha', '📦 Persediaan', '✏️ Perlengkapan', '🖥️ Peralatan']
AKUN_HUTANG = ['💳 Hutang Usaha', '🏛️ Hutang Bank', '👥 Hutang Gaji']
AKUN_MODAL = ['💎 Modal Pemilik', '💸 Prive']
AKUN_PENDAPATAN = ['💰 Pendapatan Jasa', '🛒 Pendapatan Penjualan', '✨ Pendapatan Lain-lain']
AKUN_BEBAN = ['👔 Beban Gaji', '💡 Beban Listrik', '🏠 Beban Sewa', '📋 Beban Perlengkapan', '📊 Beban Lain-lain']

SEMUA_AKUN = AKUN_HARTA + AKUN_HUTANG + AKUN_MODAL + AKUN_PENDAPATAN + AKUN_BEBAN

# Fungsi untuk menentukan jenis akun
def jenis_akun(nama_akun):
    # Hapus emoji untuk pengecekan
    nama_bersih = nama_akun.split(' ', 1)[1] if ' ' in nama_akun else nama_akun
    
    if any(nama_bersih in akun for akun in AKUN_HARTA):
        return 'Harta'
    elif any(nama_bersih in akun for akun in AKUN_HUTANG):
        return 'Hutang'
    elif any(nama_bersih in akun for akun in AKUN_MODAL):
        return 'Modal'
    elif any(nama_bersih in akun for akun in AKUN_PENDAPATAN):
        return 'Pendapatan'
    elif any(nama_bersih in akun for akun in AKUN_BEBAN):
        return 'Beban'
    return 'Lainnya'

# Fungsi untuk menghitung saldo akun
def hitung_saldo_akun(nama_akun):
    debit = 0
    kredit = 0
    for t in st.session_state.transaksi:
        if t['akun_debit'] == nama_akun:
            debit += t['jumlah']
        if t['akun_kredit'] == nama_akun:
            kredit += t['jumlah']
    
    jenis = jenis_akun(nama_akun)
    if jenis in ['Harta', 'Beban']:
        return debit - kredit
    else:
        return kredit - debit

# Fungsi untuk membuat dokumen Word
def buat_dokumen_word():
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    title = doc.add_heading('LAPORAN AKUNTANSI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tanggal_cetak = doc.add_paragraph(f'Tanggal Cetak: {datetime.now().strftime("%d-%m-%Y %H:%M")}')
    tanggal_cetak.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    
    # NERACA SALDO
    doc.add_heading('NERACA SALDO', 1)
    
    if st.session_state.transaksi:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'No'
        header_cells[1].text = 'Nama Akun'
        header_cells[2].text = 'Debit (Rp)'
        header_cells[3].text = 'Kredit (Rp)'
        
        no = 1
        total_debit = 0
        total_kredit = 0
        
        for akun in SEMUA_AKUN:
            saldo = hitung_saldo_akun(akun)
            if saldo != 0:
                jenis = jenis_akun(akun)
                if jenis in ['Harta', 'Beban']:
                    debit = abs(saldo)
                    kredit = 0
                    total_debit += debit
                else:
                    debit = 0
                    kredit = abs(saldo)
                    total_kredit += kredit
                
                row_cells = table.add_row().cells
                row_cells[0].text = str(no)
                row_cells[1].text = akun
                row_cells[2].text = f"{debit:,.0f}" if debit > 0 else "-"
                row_cells[3].text = f"{kredit:,.0f}" if kredit > 0 else "-"
                no += 1
        
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = "TOTAL"
        row_cells[2].text = f"{total_debit:,.0f}"
        row_cells[3].text = f"{total_kredit:,.0f}"
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    doc.add_page_break()
    
    # LAPORAN LABA RUGI
    doc.add_heading('LAPORAN LABA RUGI', 1)
    
    if st.session_state.transaksi:
        doc.add_heading('Pendapatan:', 2)
        table_pendapatan = doc.add_table(rows=1, cols=2)
        table_pendapatan.style = 'Light Grid Accent 1'
        
        header_cells = table_pendapatan.rows[0].cells
        header_cells[0].text = 'Akun'
        header_cells[1].text = 'Jumlah (Rp)'
        
        total_pendapatan = 0
        for akun in AKUN_PENDAPATAN:
            saldo = hitung_saldo_akun(akun)
            if saldo > 0:
                total_pendapatan += saldo
                row_cells = table_pendapatan.add_row().cells
                row_cells[0].text = akun
                row_cells[1].text = f"{saldo:,.0f}"
        
        row_cells = table_pendapatan.add_row().cells
        row_cells[0].text = "Total Pendapatan"
        row_cells[1].text = f"{total_pendapatan:,.0f}"
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        doc.add_heading('Beban:', 2)
        table_beban = doc.add_table(rows=1, cols=2)
        table_beban.style = 'Light Grid Accent 1'
        
        header_cells = table_beban.rows[0].cells
        header_cells[0].text = 'Akun'
        header_cells[1].text = 'Jumlah (Rp)'
        
        total_beban = 0
        for akun in AKUN_BEBAN:
            saldo = hitung_saldo_akun(akun)
            if saldo > 0:
                total_beban += saldo
                row_cells = table_beban.add_row().cells
                row_cells[0].text = akun
                row_cells[1].text = f"{saldo:,.0f}"
        
        row_cells = table_beban.add_row().cells
        row_cells[0].text = "Total Beban"
        row_cells[1].text = f"{total_beban:,.0f}"
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        laba_rugi = total_pendapatan - total_beban
        hasil = doc.add_paragraph()
        hasil.add_run('LABA BERSIH: ' if laba_rugi >= 0 else 'RUGI BERSIH: ').bold = True
        hasil.add_run(f'Rp {abs(laba_rugi):,.0f}').bold = True
        hasil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # DATA TRANSAKSI
    doc.add_heading('DATA TRANSAKSI', 1)
    
    if st.session_state.transaksi:
        table_transaksi = doc.add_table(rows=1, cols=5)
        table_transaksi.style = 'Light Grid Accent 1'
        
        header_cells = table_transaksi.rows[0].cells
        header_cells[0].text = 'Tanggal'
        header_cells[1].text = 'Keterangan'
        header_cells[2].text = 'Debit'
        header_cells[3].text = 'Kredit'
        header_cells[4].text = 'Jumlah (Rp)'
        
        for t in st.session_state.transaksi:
            row_cells = table_transaksi.add_row().cells
            row_cells[0].text = str(t['tanggal'])
            row_cells[1].text = t['keterangan']
            row_cells[2].text = t['akun_debit']
            row_cells[3].text = t['akun_kredit']
            row_cells[4].text = f"{t['jumlah']:,.0f}"
    
    return doc

# Header dengan styling
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.markdown("<h1>💰 Sistem Akuntansi Pro</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #7f8c8d; font-size: 16px;'>Kelola Keuangan Bisnis Anda dengan Mudah</p>", unsafe_allow_html=True)

# Tombol download di header
if st.session_state.transaksi:
    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
    
    with col2:
        if DOCX_AVAILABLE:
            doc = buat_dokumen_word()
            if doc:
                bio = io.BytesIO()
                doc.save(bio)
                st.download_button(
                    label="📄 Download Word",
                    data=bio.getvalue(),
                    file_name=f"Laporan_Akuntansi_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    with col3:
        # Dashboard Metrics
        total_transaksi = len(st.session_state.transaksi)
        st.metric("📊 Total Transaksi", total_transaksi)

st.markdown("---")

# Sidebar dengan styling
with st.sidebar:
    st.markdown("### 🎯 Menu Navigasi")
    menu = st.selectbox(
        "",
        ["🏠 Dashboard", "📝 Input Transaksi", "📖 Buku Besar", "⚖️ Neraca Saldo", "💰 Laporan Laba Rugi", "📋 Data Transaksi"],
        label_visibility="collapsed"
    )
    
    st.markdown("---")
    
    # Quick Stats di sidebar
    if st.session_state.transaksi:
        st.markdown("### 📈 Ringkasan Cepat")
        
        # Hitung total pendapatan dan beban
        total_pendapatan = sum(hitung_saldo_akun(akun) for akun in AKUN_PENDAPATAN if hitung_saldo_akun(akun) > 0)
        total_beban = sum(hitung_saldo_akun(akun) for akun in AKUN_BEBAN if hitung_saldo_akun(akun) > 0)
        laba_rugi = total_pendapatan - total_beban
        
        st.metric("💵 Pendapatan", f"Rp {total_pendapatan:,.0f}")
        st.metric("💸 Beban", f"Rp {total_beban:,.0f}")
        
        if laba_rugi >= 0:
            st.metric("✅ Laba", f"Rp {laba_rugi:,.0f}")
        else:
            st.metric("⚠️ Rugi", f"Rp {abs(laba_rugi):,.0f}")
    
    st.markdown("---")
    st.markdown("### 💡 Bantuan")
    with st.expander("📚 Panduan Cepat"):
        st.markdown("""
        **Sistem Double Entry:**
        - Setiap transaksi = Debit + Kredit
        - Total Debit = Total Kredit
        
        **Aturan Debit-Kredit:**
        - **Harta**: Debit (+) | Kredit (-)
        - **Hutang**: Debit (-) | Kredit (+)
        - **Modal**: Debit (-) | Kredit (+)
        - **Pendapatan**: Debit (-) | Kredit (+)
        - **Beban**: Debit (+) | Kredit (-)
        """)

# DASHBOARD
if menu == "🏠 Dashboard":
    st.markdown("## 🎯 Dashboard Akuntansi")
    
    if st.session_state.transaksi:
        # Metrics Row
        col1, col2, col3, col4 = st.columns(4)
        
        total_pendapatan = sum(hitung_saldo_akun(akun) for akun in AKUN_PENDAPATAN if hitung_saldo_akun(akun) > 0)
        total_beban = sum(hitung_saldo_akun(akun) for akun in AKUN_BEBAN if hitung_saldo_akun(akun) > 0)
        laba_rugi = total_pendapatan - total_beban
        total_kas = hitung_saldo_akun('💵 Kas')
        
        with col1:
            st.metric("💰 Total Pendapatan", f"Rp {total_pendapatan:,.0f}", delta="Kredit")
        with col2:
            st.metric("💸 Total Beban", f"Rp {total_beban:,.0f}", delta="Debit")
        with col3:
            if laba_rugi >= 0:
                st.metric("📈 Laba Bersih", f"Rp {laba_rugi:,.0f}", delta="Positif", delta_color="normal")
            else:
                st.metric("📉 Rugi Bersih", f"Rp {abs(laba_rugi):,.0f}", delta="Negatif", delta_color="inverse")
        with col4:
            st.metric("💵 Saldo Kas", f"Rp {total_kas:,.0f}")
        
        st.markdown("---")
        
        # Transaksi Terakhir
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📝 Transaksi Terbaru")
            transaksi_terbaru = st.session_state.transaksi[-5:][::-1]
            for t in transaksi_terbaru:
                with st.container():
                    st.markdown(f"""
                    <div style='background: white; padding: 1rem; border-radius: 10px; margin-bottom: 0.5rem; box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>
                        <strong>{t['keterangan']}</strong><br>
                        <small>{t['tanggal']} | Rp {t['jumlah']:,.0f}</small><br>
                        <small style='color: #27ae60;'>Debit: {t['akun_debit']}</small> | 
                        <small style='color: #e74c3c;'>Kredit: {t['akun_kredit']}</small>
                    </div>
                    """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("### 📊 Ringkasan Akun")
            
            # Top 5 akun dengan saldo terbesar
            saldo_akun = []
            for akun in SEMUA_AKUN:
                saldo = hitung_saldo_akun(akun)
                if saldo != 0:
                    saldo_akun.append({'Akun': akun, 'Saldo': abs(saldo)})
            
            if saldo_akun:
                saldo_akun.sort(key=lambda x: x['Saldo'], reverse=True)
                for akun in saldo_akun[:5]:
                    st.markdown(f"""
                    <div style='background: white; padding: 0.8rem; border-radius: 8px; margin-bottom: 0.5rem; box-shadow: 0 2px 4px rgba(0,0,0,0.1);'>
                        <strong>{akun['Akun']}</strong><br>
                        <span style='color: #667eea; font-size: 18px; font-weight: bold;'>Rp {akun['Saldo']:,.0f}</span>
                    </div>
                    """, unsafe_allow_html=True)
    else:
        st.info("👋 Selamat datang! Mulai dengan menambahkan transaksi pertama Anda.")
        st.markdown("""
        <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 2rem; border-radius: 15px; color: white; text-align: center;'>
            <h2>🚀 Mulai Sekarang!</h2>
            <p>Kelola keuangan bisnis Anda dengan sistem akuntansi yang mudah dan profesional</p>
        </div>
        """, unsafe_allow_html=True)

# INPUT TRANSAKSI
elif menu == "📝 Input Transaksi":
    st.markdown("## 📝 Input Transaksi Baru")
    
    with st.container():
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📅 Detail Transaksi")
            tanggal = st.date_input("Tanggal", datetime.now())
            keterangan = st.text_input("Keterangan Transaksi", placeholder="Misal: Pembayaran sewa kantor...")
            akun_debit = st.selectbox("Akun Debit", SEMUA_AKUN, help="Akun yang bertambah")
        
        with col2:
            st.markdown("### 💰 Informasi Keuangan")
            akun_kredit = st.selectbox("Akun Kredit", SEMUA_AKUN, help="Akun yang berkurang")
            jumlah = st.number_input("Jumlah (Rp)", min_value=0, step=10000, format="%d")
            st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("✅ Simpan Transaksi", type="primary", use_container_width=True):
            if jumlah > 0 and keterangan:
                transaksi_baru = {
                    'tanggal': tanggal,
                    'keterangan': keterangan,
                    'akun_debit': akun_debit,
                    'akun_kredit': akun_kredit,
                    'jumlah': jumlah
                }
                st.session_state.transaksi.append(transaksi_baru)
                st.success("✅ Transaksi berhasil ditambahkan!")
                st.balloons()
            else:
                st.error("❌ Mohon lengkapi semua field!")
    
    # Contoh transaksi dengan card design
    st.markdown("---")
    st.markdown("### 💡 Contoh Transaksi")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 15px; color: white;'>
            <h4>📥 Setoran Modal</h4>
            <p><strong>Nominal:</strong> Rp 10.000.000</p>
            <p><strong>Debit:</strong> 💵 Kas</p>
            <p><strong>Kredit:</strong> 💎 Modal Pemilik</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 1.5rem; border-radius: 15px; color: white;'>
            <h4>💸 Bayar Beban</h4>
            <p><strong>Nominal:</strong> Rp 2.000.000</p>
            <p><strong>Debit:</strong> 🏠 Beban Sewa</p>
            <p><strong>Kredit:</strong> 💵 Kas</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style='background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 1.5rem; border-radius: 15px; color: white;'>
            <h4>💰 Terima Pendapatan</h4>
            <p><strong>Nominal:</strong> Rp 5.000.000</p>
            <p><strong>Debit:</strong> 💵 Kas</p>
            <p><strong>Kredit:</strong> 💰 Pendapatan Jasa</p>
        </div>
        """, unsafe_allow_html=True)

# BUKU BESAR
elif menu == "📖 Buku Besar":
    st.markdown("## 📖 Buku Besar")
    
    akun_dipilih = st.selectbox("🔍 Pilih Akun", SEMUA_AKUN)
    
    transaksi_akun = []
    saldo = 0
    
    for t in st.session_state.transaksi:
        if t['akun_debit'] == akun_dipilih or t['akun_kredit'] == akun_dipilih:
            debit = t['jumlah'] if t['akun_debit'] == akun_dipilih else 0
            kredit = t['jumlah'] if t['akun_kredit'] == akun_dipilih else 0
            
            jenis = jenis_akun(akun_dipilih)
            if jenis in ['Harta', 'Beban']:
                saldo += debit - kredit
            else:
                saldo += kredit - debit
            
            transaksi_akun.append({
                'Tanggal': t['tanggal'],
                'Keterangan': t['keterangan'],
                'Debit': f"Rp {debit:,.0f}" if debit > 0 else "-",
                'Kredit': f"Rp {kredit:,.0f}" if kredit > 0 else "-",
                'Saldo': f"Rp {saldo:,.0f}"
            })
    
    if transaksi_akun:
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            st.markdown(f"### {akun_dipilih}")
        with col3:
            st.metric("💰 Saldo Akhir", f"Rp {saldo:,.0f}")
        
        df = pd.DataFrame(transaksi_akun)
        st.dataframe(df, use_container_width=True, height=400)
        
        if EXCEL_AVAILABLE:
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Buku Besar', index=False)
            
            st.download_button(
                label="📥 Download Excel",
                data=output.getvalue(),
                file_name=f"Buku_Besar_{akun_dipilih}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    else:
        st.info("📭 Belum ada transaksi untuk akun ini")

# NERACA SALDO
elif menu == "⚖️ Neraca Saldo":
    st.markdown("## ⚖️ Neraca Saldo")
    
    if st.session_state.transaksi:
        data_neraca = []
        total_debit = 0
        total_kredit = 0
        
        for akun in SEMUA_AKUN:
            saldo = hitung_saldo_akun(akun)
            if saldo != 0:
                jenis = jenis_akun(akun)
                if jenis in ['Harta', 'Beban']:
                    debit = abs(saldo)
                    kredit = 0
                    total_debit += debit
                else:
                    debit = 0
                    kredit = abs(saldo)
                    total_kredit += kredit
                
                data_neraca.append({
                    'Akun': akun,
                    'Jenis': jenis,
                    'Debit': f"Rp {debit:,.0f}" if debit > 0 else "-",
                    'Kredit': f"Rp {kredit:,.0f}" if kredit > 0 else "-"
                })
        
        if data_neraca:
            df_neraca = pd.DataFrame(data_neraca)
            st.dataframe(df_neraca, use_container_width=True, height=400)
            
            st.markdown("---")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📊 Total Debit", f"Rp {total_debit:,.0f}")
            with col2:
                st.metric("📊 Total Kredit", f"Rp {total_kredit:,.0f}")
            with col3:
                if EXCEL_AVAILABLE:
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        df_neraca.to_excel(writer, sheet_name='Neraca Saldo', index=False)
                    
                    st.download_button(
                        label="📥 Download Excel",
                        data=output.getvalue(),
                        file_name=f"Neraca_Saldo_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            if total_debit == total_kredit:
                st.success("✅ Neraca Saldo Seimbang! Pembukuan Anda Benar.")
            else:
                st.error(f"❌ Neraca Tidak Seimbang! Selisih: Rp {abs(total_debit - total_kredit):,.0f}")
        else:
            st.info("📭 Belum ada data untuk ditampilkan")
    else:
        st.info("📭 Belum ada transaksi. Silakan input transaksi terlebih dahulu.")

# LAPORAN LABA RUGI
elif menu == "💰 Laporan Laba Rugi":
    st.markdown("## 💰 Laporan Laba Rugi")
    
    if st.session_state.transaksi:
        # Hitung total pendapatan
        total_pendapatan = 0
        data_pendapatan = []
        for akun in AKUN_PENDAPATAN:
            saldo = hitung_saldo_akun(akun)
            if saldo > 0:
                total_pendapatan += saldo
                data_pendapatan.append({'Akun': akun, 'Jumlah': f"Rp {saldo:,.0f}"})
        
        # Hitung total beban
        total_beban = 0
        data_beban = []
        for akun in AKUN_BEBAN:
            saldo = hitung_saldo_akun(akun)
            if saldo > 0:
                total_beban += saldo
                data_beban.append({'Akun': akun, 'Jumlah': f"Rp {saldo:,.0f}"})
        
        col1, col2 = st.columns(2)
        
        # Tampilkan Pendapatan
        with col1:
            st.markdown("### 📈 Pendapatan")
            if data_pendapatan:
                df_pendapatan = pd.DataFrame(data_pendapatan)
                st.dataframe(df_pendapatan, use_container_width=True, hide_index=True)
            st.metric("Total Pendapatan", f"Rp {total_pendapatan:,.0f}")
        
        # Tampilkan Beban
        with col2:
            st.markdown("### 📉 Beban")
            if data_beban:
                df_beban = pd.DataFrame(data_beban)
                st.dataframe(df_beban, use_container_width=True, hide_index=True)
            st.metric("Total Beban", f"Rp {total_beban:,.0f}")
        
        st.markdown("---")
        
        # Hitung laba/rugi
        laba_rugi = total_pendapatan - total_beban
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if laba_rugi > 0:
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%); padding: 2rem; border-radius: 15px; text-align: center; color: white;'>
                    <h2>🎉 LABA BERSIH</h2>
                    <h1 style='font-size: 3rem; margin: 0;'>Rp {laba_rugi:,.0f}</h1>
                    <p style='margin-top: 1rem; opacity: 0.9;'>Selamat! Bisnis Anda menguntungkan</p>
                </div>
                """, unsafe_allow_html=True)
            elif laba_rugi < 0:
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, #eb3349 0%, #f45c43 100%); padding: 2rem; border-radius: 15px; text-align: center; color: white;'>
                    <h2>📛 RUGI BERSIH</h2>
                    <h1 style='font-size: 3rem; margin: 0;'>Rp {abs(laba_rugi):,.0f}</h1>
                    <p style='margin-top: 1rem; opacity: 0.9;'>Perlu evaluasi strategi bisnis</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 2rem; border-radius: 15px; text-align: center; color: white;'>
                    <h2>⚖️ IMPAS (BREAK EVEN)</h2>
                    <h1 style='font-size: 3rem; margin: 0;'>Rp 0</h1>
                    <p style='margin-top: 1rem; opacity: 0.9;'>Pendapatan = Beban</p>
                </div>
                """, unsafe_allow_html=True)
        
        # Download Excel Laba Rugi
        if EXCEL_AVAILABLE:
            st.markdown("---")
            col1, col2, col3 = st.columns([1, 1, 1])
            with col2:
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    if data_pendapatan:
                        df_pend = pd.DataFrame(data_pendapatan)
                        df_pend.to_excel(writer, sheet_name='Pendapatan', index=False)
                    if data_beban:
                        df_beb = pd.DataFrame(data_beban)
                        df_beb.to_excel(writer, sheet_name='Beban', index=False)
                    
                    df_ringkasan = pd.DataFrame({
                        'Keterangan': ['Total Pendapatan', 'Total Beban', 'Laba/Rugi Bersih'],
                        'Jumlah': [f"Rp {total_pendapatan:,.0f}", f"Rp {total_beban:,.0f}", f"Rp {laba_rugi:,.0f}"]
                    })
                    df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
                
                st.download_button(
                    label="📥 Download Excel Laba Rugi",
                    data=output.getvalue(),
                    file_name=f"Laporan_Laba_Rugi_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
    else:
        st.info("📭 Belum ada transaksi. Silakan input transaksi terlebih dahulu.")

# DATA TRANSAKSI
elif menu == "📋 Data Transaksi":
    st.markdown("## 📋 Data Transaksi")
    
    if st.session_state.transaksi:
        df_transaksi = pd.DataFrame(st.session_state.transaksi)
        
        # Format tampilan
        df_display = df_transaksi.copy()
        df_display['jumlah'] = df_display['jumlah'].apply(lambda x: f"Rp {x:,.0f}")
        df_display.columns = ['Tanggal', 'Keterangan', 'Debit', 'Kredit', 'Jumlah']
        
        st.dataframe(df_display, use_container_width=True, height=400)
        
        st.markdown("---")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("📊 Total Transaksi", len(st.session_state.transaksi))
        
        with col2:
            if EXCEL_AVAILABLE:
                output = io.BytesIO()
                df_transaksi.to_excel(output, index=False, engine='openpyxl')
                
                st.download_button(
                    label="📥 Download Excel",
                    data=output.getvalue(),
                    file_name=f"Data_Transaksi_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
        
        with col3:
            if st.button("🗑️ Hapus Semua Data", type="secondary", use_container_width=True):
                if st.session_state.get('confirm_delete', False):
                    st.session_state.transaksi = []
                    st.session_state.confirm_delete = False
                    st.rerun()
                else:
                    st.session_state.confirm_delete = True
                    st.warning("⚠️ Klik sekali lagi untuk konfirmasi hapus!")
    else:
        st.info("📭 Belum ada transaksi yang tercatat")
        st.markdown("""
        <div style='text-align: center; padding: 3rem;'>
            <h2>🚀 Mulai Input Transaksi</h2>
            <p style='color: #7f8c8d;'>Gunakan menu "Input Transaksi" untuk memulai</p>
        </div>
        """, unsafe_allow_html=True)

# Footer
st.markdown("---")
col1, col2, col3 = st.columns(3)
with col2:
    st.markdown("""
    <div style='text-align: center; color: #7f8c8d; padding: 1rem;'>
        <p>💼 Sistem Akuntansi Pro | Dibuat dengan  menggunakan Streamlit</p>
        <small>© 2025 - Kelola keuangan bisnis Anda dengan mudah</small>
    </div>
    """, unsafe_allow_html=True)